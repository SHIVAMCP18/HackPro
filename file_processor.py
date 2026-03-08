"""
file_processor.py — High-performance PII detection and sanitization
Optimized for 20MB+ files:
  - Single-pass hint-gated regex substitution
  - Parallel processing for PDF pages and DOCX paragraphs
  - Chunked streaming for large TXT/SQL files
  - Column-batch CSV processing
  - Pre-compiled patterns at module load

Fix: DOCX embedded-image redaction now uses a proper two-step ZIP rebuild
     (save doc first → reopen ZIP → write new ZIP → close → getvalue)
     so the ZIP central directory is fully written before bytes are read.
"""
import io
import zipfile
import csv
import re
from concurrent.futures import ThreadPoolExecutor, as_completed

import pdfplumber
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from PIL import Image, ImageDraw

from pii_engine import build_pii_summary, REGEX_PATTERNS, name_address_scan

# ── CONFIG ────────────────────────────────────────────────────────
MAX_WORKERS = 8        # parallel threads for PDF/DOCX
CHUNK_SIZE  = 500_000  # bytes per chunk for TXT/SQL streaming
OVERLAP     = 200      # overlap between chunks to avoid splitting PII at boundary

# ── PRE-COMPILE EVERYTHING ONCE AT MODULE LOAD ───────────────────
def _clean_pattern(p):
    return re.sub(r'\(\?[a-zA-Z]+\)', '', p)

# Individual compiled patterns — used with hint-based skipping
_COMPILED_EACH = {
    pt: (re.compile(_clean_pattern(p), re.IGNORECASE), m)
    for pt, (p, m) in REGEX_PATTERNS.items()
}

# Hint functions — fast str/re checks to skip patterns that can't match
_HINTS = {
    "aadhaar":        lambda t: bool(re.search(r'\d{4} \d{4}', t[:500] if len(t) > 500 else t)),
    "us_phone":       lambda t: "(" in t or "+1" in t,
    "email":          lambda t: "@" in t,
    "ip_address":     lambda t: t.count(".") > 5,
    "passport":       lambda t: bool(re.search(r'[A-Z]\d{7}', t[:500] if len(t) > 500 else t)),
    "ifsc":           lambda t: bool(re.search(r'[A-Z]{4}0', t[:500] if len(t) > 500 else t)),
    "upi":            lambda t: "@" in t and any(x in t.lower() for x in ("upi", "ybl", "okaxis", "paytm")),
    "credit_card":    lambda t: bool(re.search(r'\d{4}[ -]\d{4}', t[:500] if len(t) > 500 else t)),
    "cvv":            lambda t: "cvv" in t.lower(),
    "expiry_date":    lambda t: "/" in t,
    "device_id":      lambda t: "android" in t.lower() or "ios" in t.lower(),
    "fingerprint":    lambda t: "fp_hash" in t,
    "face_template":  lambda t: "face_tmp" in t,
    "dob":            lambda t: bool(re.search(r'\d{2}[/-]\d{2}[/-]\d{4}', t[:500] if len(t) > 500 else t)),
    "pincode":        lambda t: "pin" in t.lower() or "zip" in t.lower(),
    "vehicle_number": lambda t: bool(re.search(r'\b[A-Z]{2}\d{2}[A-Z]', t[:500] if len(t) > 500 else t)),
    "voter_id":       lambda t: bool(re.search(r'\b[A-Z]{3}\d{7}', t[:500] if len(t) > 500 else t)),
    "gstin":          lambda t: bool(re.search(r'\b\d{2}[A-Z]{5}\d{4}', t[:500] if len(t) > 500 else t)),
    "swift_code":     lambda t: "swift" in t.lower() or "bic" in t.lower() or "iban" in t.lower(),
    "account_number": lambda t: bool(re.search(r'\d{9,18}', t[:500] if len(t) > 500 else t)),
}

# Combined regex — still used for image OCR token matching
_PATTERN_LIST = list(REGEX_PATTERNS.items())
_GROUP_MAP    = {}
_parts        = []
for idx, (pii_type, (pattern, masker)) in enumerate(_PATTERN_LIST):
    group_name = f"pii{idx}"
    clean = _clean_pattern(pattern)
    _parts.append(f"(?P<{group_name}>{clean})")
    _GROUP_MAP[group_name] = (pii_type, masker)
_COMBINED = re.compile("|".join(_parts), re.IGNORECASE)

# Name/address patterns pre-compiled
_NAME_LABELED = re.compile(
    r'(?i)(?:name|patient|employee|customer|client|user|applicant|member)[\s:.-]+'
    r'((?:[A-Z][a-z]+)(?:\s+(?:[A-Z][a-z]+))*)(?=\s+[A-Z]|\s*$|[,\n])'
)
_NAME_TITLED = re.compile(
    r'(?:Mr\.?|Mrs\.?|Ms\.?|Dr\.?|Prof\.?|Shri|Smt\.?|Sri|Er\.?)\s+'
    r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2})'
)
_ADDRESS = re.compile(
    r'\b\d+[,\s]+[A-Za-z0-9\s,\-\.]+(?:Road|Street|Nagar|Colony|Sector|Phase|Block|'
    r'Avenue|Lane|Marg|Chowk|Layout|Extension|Society|Residency|Apartments?|Towers?|'
    r'Floor|Flat|Plot)[^.\n]{0,60}',
    re.IGNORECASE
)


# ── CORE: single-pass hint-gated scan ────────────────────────────

def _single_pass_scan(text: str) -> tuple:
    """
    Hint-gated per-pattern scan — skips patterns that cannot match via
    cheap string checks before running expensive regex on large files.
    ~5x faster than combined mega-pattern on 10MB+ files.
    """
    detections = []
    masked = text

    for pii_type, (pat, masker) in _COMPILED_EACH.items():
        hint = _HINTS.get(pii_type)
        if hint and not hint(masked):
            continue
        found = []

        def _replace(m, _pt=pii_type, _mk=masker, _f=found):
            v = m.group()
            try:
                mv = _mk(v)
            except Exception:
                mv = "[REDACTED]"
            _f.append({"pii_type": _pt, "original_value": v,
                       "masked_value": mv, "detection_method": "regex", "confidence": 1.0})
            return mv

        masked = pat.sub(_replace, masked)
        detections.extend(found)

    # ── Name pass (4 complementary patterns) ─────────────────────
    from pii_engine import (NAME_PATTERN, NAME_TITLED, NAME_VERB_PREFIX,
                             NAME_TITLE_PAIR, _NAME_COMMON_WORDS)

    def _add_name(masked, name_val, method, conf=0.9):
        name_val = name_val.strip()
        if len(name_val) > 2 and "[NAME REDACTED]" not in name_val:
            detections.append({"pii_type": "name", "original_value": name_val,
                               "masked_value": "[NAME REDACTED]",
                               "detection_method": method, "confidence": conf})
            masked = masked.replace(name_val, "[NAME REDACTED]", 1)
        return masked

    # 1. Labeled
    if any(kw in masked for kw in ("Name:", "name:", "customer", "Customer",
                                    "patient", "employee", "client", "member")):
        for m in NAME_PATTERN.finditer(masked):
            masked = _add_name(masked, m.group(1), "pattern-labeled")
    # 2. Titled
    for m in NAME_TITLED.finditer(masked):
        masked = _add_name(masked, m.group(1), "pattern-titled")
    # 3. Verb-prefixed
    for m in NAME_VERB_PREFIX.finditer(masked):
        masked = _add_name(masked, m.group(1), "pattern-verb", conf=0.88)
    # 4. Standalone Title-Case pair
    for m in NAME_TITLE_PAIR.finditer(masked):
        first, last = m.group(1), m.group(2)
        if first not in _NAME_COMMON_WORDS and last not in _NAME_COMMON_WORDS:
            masked = _add_name(masked, f"{first} {last}", "pattern-titlecase", conf=0.82)

    # ── Address pass ──────────────────────────────────────────────
    from pii_engine import ADDRESS_PATTERN as _ADDR_PAT
    if any(kw in masked for kw in ("Road", "Street", "Nagar", "Colony", "Sector", "Phase", "Block")):
        for match in _ADDR_PAT.finditer(masked):
            addr = match.group().strip()
            if len(addr) > 10:
                parts = [p.strip() for p in addr.split(",")]
                num_match = re.match(r'(\d+[A-Za-z]?)', parts[0].strip())
                masked_addr = (f"{num_match.group(1)}, ***, ***, ***" if num_match else "***, ***, ***")
                detections.append({"pii_type": "address", "original_value": addr,
                                   "masked_value": masked_addr, "detection_method": "pattern", "confidence": 0.85})
                masked = masked.replace(addr, masked_addr, 1)

    return masked, detections


def _fast_scan(text: str) -> tuple:
    """Public alias used by all processors."""
    return _single_pass_scan(text)


# ── CHUNKED SCAN for 20MB+ text files ────────────────────────────

def _chunked_scan(text: str) -> tuple:
    """
    Split large text into overlapping chunks, process in parallel,
    then stitch results. Prevents memory spikes on 20MB+ files.
    """
    if len(text) <= CHUNK_SIZE:
        return _single_pass_scan(text)

    chunks = []
    starts = []
    i = 0
    while i < len(text):
        end = min(i + CHUNK_SIZE, len(text))
        if end < len(text):
            newline = text.rfind('\n', i + CHUNK_SIZE - OVERLAP, end)
            if newline > i:
                end = newline + 1
        chunks.append(text[i:end])
        starts.append(i)
        i = end - OVERLAP if end < len(text) else end

    results = [None] * len(chunks)
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as ex:
        futures = {ex.submit(_single_pass_scan, chunk): idx
                   for idx, chunk in enumerate(chunks)}
        for future in as_completed(futures):
            idx = futures[future]
            results[idx] = future.result()

    all_detections = []
    stitched = []
    for idx, (masked_chunk, dets) in enumerate(results):
        all_detections.extend(dets)
        if idx < len(results) - 1:
            stitched.append(masked_chunk[:-OVERLAP] if len(masked_chunk) > OVERLAP else masked_chunk)
        else:
            stitched.append(masked_chunk)

    return "".join(stitched), all_detections


# ── ROUTER ───────────────────────────────────────────────────────

def process_file(file_bytes: bytes, filename: str) -> tuple:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return _process_pdf(file_bytes)
    elif ext == "docx":
        return _process_docx(file_bytes)
    elif ext in ("sql", "txt"):
        return _process_text(file_bytes)
    elif ext == "csv":
        return _process_csv(file_bytes)
    elif ext == "json":
        return _process_json(file_bytes)
    elif ext in ("png", "jpg", "jpeg"):
        return _process_image(file_bytes, ext)
    else:
        return _process_text(file_bytes)


# ── IMAGE ─────────────────────────────────────────────────────────

_nlp = None

def _get_nlp():
    global _nlp
    if _nlp is None:
        try:
            import spacy
            _nlp = spacy.load("en_core_web_sm")
        except Exception:
            _nlp = False
    return _nlp if _nlp else None


def _process_image(file_bytes: bytes, ext: str = "png") -> tuple:
    try:
        import pytesseract
        from pytesseract import Output
        from PIL import ImageFilter, ImageEnhance

        image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        orig_w, orig_h = image.size

        # 2x upscale + contrast + sharpen → dramatically improves OCR accuracy
        ocr_img = image.resize((orig_w * 2, orig_h * 2), Image.LANCZOS).convert("L")
        ocr_img = ImageEnhance.Contrast(ocr_img).enhance(2.0)
        ocr_img = ocr_img.filter(ImageFilter.SHARPEN)

        scale_x, scale_y = 0.5, 0.5
        draw = ImageDraw.Draw(image)
        d = pytesseract.image_to_data(ocr_img, output_type=Output.DICT, config="--psm 6")
        n = len(d["text"])

        full_text = " ".join(word for word in d["text"] if word.strip())

        def _norm(t):
            t = re.sub(r'\b(\d{4})(\d{4})(\d{4})\b', r'\1 \2 \3', t)
            t = re.sub(r'\b(\d{4})(\d{4})\b', r'\1 \2', t)
            return t

        norm_text = _norm(full_text)

        # pii_substrings: all lowercased token variants of detected PII values.
        # We store both the space-split tokens AND the fully-merged (no-space) form
        # so we can match regardless of whether OCR returned "1234 5678 9012" or
        # "123456789012" for the same Aadhaar number.
        pii_substrings = set()

        _LABEL_WORDS = {
            'name', 'pan', 'phone', 'aadhaar', 'aadhar', 'email', 'dob',
            'address', 'account', 'ifsc', 'upi', 'gstin', 'passport',
            'voter', 'vehicle', 'cvv', 'card', 'patient', 'employee',
            'customer', 'client', 'user', 'mobile', 'contact', 'no',
            'number', 'id', 'date', 'birth', 'pincode', 'zip', 'swift',
            'by', 'the', 'and', 'for', 'ref', 'to', 'of', 'in', 'at',
        }

        from pii_engine import (NAME_PATTERN as _NP, NAME_TITLED as _NT,
                                 NAME_VERB_PREFIX as _NVP, NAME_TITLE_PAIR as _NTP,
                                 _NAME_COMMON_WORDS as _NCW)

        def _add_pii_value(val):
            """Register all token variants of a PII value into pii_substrings."""
            clean_val = re.sub(r'^[^a-zA-Z0-9@+]|[^a-zA-Z0-9@.]$', '', val)
            # 1. Fully merged (no spaces) — catches OCR that runs digits together
            merged = re.sub(r'\s+', '', clean_val).lower()
            if merged and len(merged) > 2:
                pii_substrings.add(merged)
                # 2. For purely numeric values (e.g. Aadhaar "1234 5678 9012"),
                #    also store ALL contiguous sub-groups OCR might split on:
                #    4+8, 8+4, 4+4+4, etc.  We do this by sliding a window over
                #    the digit string and storing every prefix/suffix of length ≥ 4.
                if merged.isdigit():
                    digs = merged
                    for start in range(0, len(digs)):
                        for end in range(start + 4, len(digs) + 1):
                            pii_substrings.add(digs[start:end])
            # 3. Per-whitespace-token — handles OCR that does split on spaces
            for token in clean_val.split():
                clean = re.sub(r'^[^a-zA-Z0-9@+]|[^a-zA-Z0-9@.]$', '', token).lower()
                if clean and clean not in _LABEL_WORDS and len(clean) > 2:
                    pii_substrings.add(clean)

        for match in _COMBINED.finditer(norm_text):
            for group_name, (pii_type, masker) in _GROUP_MAP.items():
                val = match.group(group_name)
                if val is not None:
                    _add_pii_value(val)
                    break

        for match in _NP.finditer(full_text):
            name_val = match.group(1)
            if name_val:
                _add_pii_value(name_val.strip())

        for match in _NT.finditer(full_text):
            _add_pii_value(match.group(1))

        for match in _NVP.finditer(full_text):
            _add_pii_value(match.group(1))

        for match in _NTP.finditer(full_text):
            first, last = match.group(1), match.group(2)
            if first not in _NCW and last not in _NCW:
                _add_pii_value(f"{first} {last}")

        nlp = _get_nlp()
        if nlp:
            for ent in nlp(full_text).ents:
                if ent.label_ in ["PERSON", "GPE"]:
                    _add_pii_value(ent.text)

        all_detections = []
        for match in _COMBINED.finditer(norm_text):
            for group_name, (pii_type, masker) in _GROUP_MAP.items():
                val = match.group(group_name)
                if val is not None:
                    try:
                        masked_val = masker(val)
                    except Exception:
                        masked_val = "****"
                    all_detections.append({
                        "pii_type": pii_type, "original_value": val,
                        "masked_value": masked_val,
                        "detection_method": "regex+ocr", "confidence": 1.0
                    })
                    break

        for match in _NP.finditer(full_text):
            name_val = match.group(1)
            if name_val:
                all_detections.append({"pii_type": "name", "original_value": name_val.strip(),
                                       "masked_value": "[NAME REDACTED]",
                                       "detection_method": "pattern+ocr", "confidence": 0.9})
        for match in _NT.finditer(full_text):
            all_detections.append({"pii_type": "name", "original_value": match.group(1).strip(),
                                   "masked_value": "[NAME REDACTED]",
                                   "detection_method": "pattern+ocr", "confidence": 0.9})
        for match in _NVP.finditer(full_text):
            all_detections.append({"pii_type": "name", "original_value": match.group(1).strip(),
                                   "masked_value": "[NAME REDACTED]",
                                   "detection_method": "pattern+ocr", "confidence": 0.88})
        for match in _NTP.finditer(full_text):
            first, last = match.group(1), match.group(2)
            if first not in _NCW and last not in _NCW:
                all_detections.append({"pii_type": "name", "original_value": f"{first} {last}",
                                       "masked_value": "[NAME REDACTED]",
                                       "detection_method": "pattern+ocr", "confidence": 0.82})

        # Draw black boxes over detected PII tokens.
        # We match each OCR word in two ways:
        #   1. Exact match: clean_word is directly in pii_substrings
        #   2. Merged match: strip all spaces from clean_word and check again
        #      (catches cases where OCR returned "123456789012" but PII was
        #       detected as "1234 5678 9012" whose merged form "123456789012" is stored)
        for i in range(n):
            word = d["text"][i].strip()
            if not word:
                continue
            clean_word = re.sub(r'^[^a-zA-Z0-9@+]|[^a-zA-Z0-9@.]$', '', word).lower()
            merged_word = re.sub(r'\s+', '', clean_word)
            if clean_word in pii_substrings or merged_word in pii_substrings:
                x, y, ww, hh = d["left"][i], d["top"][i], d["width"][i], d["height"][i]
                if ww > 0 and hh > 0:
                    sx  = int(x  * scale_x)
                    sy  = int(y  * scale_y)
                    sww = int(ww * scale_x)
                    shh = int(hh * scale_y)
                    draw.rectangle([sx - 2, sy - 2, sx + sww + 2, sy + shh + 2], fill="black")

        out = io.BytesIO()
        image.save(out, format="JPEG" if ext in ("jpg", "jpeg") else "PNG")
        return out.getvalue(), all_detections, build_pii_summary(all_detections)

    except ImportError:
        image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        draw = ImageDraw.Draw(image)
        iw, ih = image.size
        draw.rectangle([0, ih // 2 - 40, iw, ih // 2 + 40], fill=(150, 0, 0))
        draw.text((20, ih // 2 - 15), "Install tesseract for OCR redaction", fill="white")
        out = io.BytesIO()
        image.save(out, format="PNG")
        return out.getvalue(), [], {}
    except Exception as e:
        print(f"[Image error] {e}")
        return file_bytes, [], {}


# ── PDF ───────────────────────────────────────────────────────────

def _process_pdf(file_bytes: bytes) -> tuple:
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_texts = [page.extract_text() or "" for page in pdf.pages]

    results = [None] * len(page_texts)
    with ThreadPoolExecutor(max_workers=min(MAX_WORKERS, max(1, len(page_texts)))) as ex:
        futures = {ex.submit(_single_pass_scan, txt): i for i, txt in enumerate(page_texts)}
        for future in as_completed(futures):
            results[futures[future]] = future.result()

    all_detections = []
    full_masked_text = ""
    for masked, dets in results:
        all_detections.extend(dets)
        full_masked_text += masked + "\n\n"

    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=A4)
    styles = getSampleStyleSheet()
    story = [
        Paragraph(
            line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"),
            styles["Normal"]
        )
        for line in full_masked_text.split("\n") if line.strip()
    ]
    doc.build(story)
    return out.getvalue(), all_detections, build_pii_summary(all_detections)


# ── DOCX ──────────────────────────────────────────────────────────

def _redact_images_in_docx_zip(docx_bytes: bytes) -> tuple:
    """
    Open the DOCX ZIP, OCR-redact every image in word/media/,
    write a new ZIP with the redacted images, and return its bytes.

    BUG FIX: The output BytesIO buffer is only valid AFTER the ZipFile
    is fully closed (so its central directory is flushed).  We therefore
    close the writer first, then call getvalue() — never inside the
    'with' block.
    """
    IMG_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.webp'}
    all_detections = []

    # --- read phase ---
    with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as src_zip:
        items = src_zip.infolist()
        file_data = {item.filename: src_zip.read(item.filename) for item in items}

    # --- redact images ---
    redacted_data = {}
    for filename, data in file_data.items():
        fname_lower = filename.lower()
        if fname_lower.startswith('word/media/') and any(fname_lower.endswith(e) for e in IMG_EXTS):
            ext = fname_lower.rsplit('.', 1)[-1]
            try:
                redacted_bytes, detections, _ = _process_image(data, ext)
                all_detections.extend(detections)
                redacted_data[filename] = redacted_bytes
            except Exception as e:
                print(f"[DOCX image error] {filename}: {e}")
                redacted_data[filename] = data  # keep original on failure
        else:
            redacted_data[filename] = data

    # --- write phase: close BEFORE getvalue() ---
    out_buf = io.BytesIO()
    with zipfile.ZipFile(out_buf, 'w', compression=zipfile.ZIP_DEFLATED) as out_zip:
        for item in items:
            out_zip.writestr(item, redacted_data[item.filename])
    # out_zip is now closed → central directory is written → getvalue() is safe
    return out_buf.getvalue(), all_detections


def _process_docx(file_bytes: bytes) -> tuple:
    all_detections = []

    # ── Step 1: Redact embedded images via ZIP manipulation ───────
    try:
        img_docx_bytes, img_detections = _redact_images_in_docx_zip(file_bytes)
        all_detections.extend(img_detections)
        file_bytes = img_docx_bytes  # continue with image-redacted bytes
    except Exception as e:
        print(f"[DOCX image pass error] {e}")

    # ── Step 2: Redact text in paragraphs and tables in parallel ──
    doc = Document(io.BytesIO(file_bytes))

    para_list = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                para_list.extend(cell.paragraphs)

    texts = ["".join(r.text for r in p.runs) for p in para_list]
    non_empty = [(i, t) for i, t in enumerate(texts) if t.strip()]

    if non_empty:
        results = {}
        with ThreadPoolExecutor(max_workers=min(MAX_WORKERS, len(non_empty))) as ex:
            futures = {ex.submit(_single_pass_scan, t): i for i, t in non_empty}
            for future in as_completed(futures):
                i = futures[future]
                results[i] = future.result()

        for i, _ in non_empty:
            masked, dets = results[i]
            all_detections.extend(dets)
            para = para_list[i]
            if para.runs:
                para.runs[0].text = masked
                for r in para.runs[1:]:
                    r.text = ""

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), all_detections, build_pii_summary(all_detections)


# ── TXT / SQL ─────────────────────────────────────────────────────

def _process_text(file_bytes: bytes) -> tuple:
    text = file_bytes.decode("utf-8", errors="replace")
    masked, detections = _chunked_scan(text)
    return masked.encode("utf-8"), detections, build_pii_summary(detections)


# ── CSV ───────────────────────────────────────────────────────────

def _process_csv(file_bytes: bytes) -> tuple:
    text = file_bytes.decode("utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    if not rows:
        return file_bytes, [], {}

    num_cols = max(len(r) for r in rows)
    sanitized = [list(r) for r in rows]
    all_detections = []
    DELIM = "\x00SPLIT\x00"

    def _process_column(col_idx):
        col_vals, row_indices = [], []
        for row_idx, row in enumerate(rows):
            if col_idx < len(row) and row[col_idx].strip():
                col_vals.append(row[col_idx])
                row_indices.append(row_idx)
        if not col_vals:
            return col_idx, [], []
        joined = DELIM.join(col_vals)
        masked_joined, dets = _single_pass_scan(joined)
        return col_idx, dets, list(zip(row_indices, masked_joined.split(DELIM)))

    with ThreadPoolExecutor(max_workers=min(MAX_WORKERS, num_cols)) as ex:
        futures = [ex.submit(_process_column, c) for c in range(num_cols)]
        for future in as_completed(futures):
            col_idx, dets, row_updates = future.result()
            all_detections.extend(dets)
            for row_idx, val in row_updates:
                sanitized[row_idx][col_idx] = val

    out = io.StringIO()
    csv.writer(out).writerows(sanitized)
    return out.getvalue().encode("utf-8"), all_detections, build_pii_summary(all_detections)


# ── JSON ──────────────────────────────────────────────────────────

def _process_json(file_bytes: bytes) -> tuple:
    text = file_bytes.decode("utf-8", errors="replace")
    masked, detections = _chunked_scan(text)
    return masked.encode("utf-8"), detections, build_pii_summary(detections)


# ── PREVIEW ───────────────────────────────────────────────────────

def extract_preview_text(file_bytes: bytes, filename: str, max_chars: int = 2000) -> str:
    """Extract a short preview. For large files, reads only the first page/chunk."""
    ext = filename.rsplit(".", 1)[-1].lower()
    try:
        if ext == "pdf":
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text = ""
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n"
                    if len(text) >= max_chars:
                        break  # stop after we have enough — don't parse all pages
                return text[:max_chars]
        elif ext == "docx":
            doc = Document(io.BytesIO(file_bytes))
            text = ""
            for p in doc.paragraphs:
                text += p.text + "\n"
                if len(text) >= max_chars:
                    break
            return text[:max_chars]
        else:
            # For text-based files just slice the raw bytes — very fast
            return file_bytes[:max_chars * 4].decode("utf-8", errors="replace")[:max_chars]
    except Exception:
        return "[Preview unavailable]"
